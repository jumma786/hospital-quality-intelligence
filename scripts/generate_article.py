"""
generate_article.py
--------------------
Builds a Medium-style article about the Hospital Quality Intelligence project
as a formatted Word document (.docx), with embedded charts.

Run:
    python scripts/generate_article.py
Output:
    docs/Hospital_Quality_Intelligence_Article.docx
"""

from __future__ import annotations
import pathlib

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches

ROOT = pathlib.Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
OUT_DIR.mkdir(exist_ok=True)
IMG = ROOT / "ml" / "outputs"
REPO = "https://github.com/jumma786/hospital-quality-intelligence"
DEMO = "https://hospital-quality-intelligence-puoehkx7u7po2laff5rv6y.streamlit.app/"

ACCENT = RGBColor(0x25, 0x63, 0xEB)
GREY = RGBColor(0x66, 0x66, 0x66)


def h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
    p.space_after = Pt(6)


def subtitle(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(13)
    r.font.color.rgb = GREY


def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = ACCENT
    p.space_before = Pt(14)
    p.space_after = Pt(4)


def body(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11.5)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.3
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.size = Pt(11.5)


def pull_quote(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = ACCENT
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)


def image(doc, path, caption):
    if not path.exists():
        return
    doc.add_picture(str(path), width=Inches(6.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = GREY


def kv_table(doc, rows):
    """Two-column label/value table (left column bold)."""
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Light List Accent 1"
    for i, (label, value) in enumerate(rows):
        c0, c1 = t.rows[i].cells
        r0 = c0.paragraphs[0].add_run(label)
        r0.bold = True
        r0.font.size = Pt(10.5)
        r1 = c1.paragraphs[0].add_run(value)
        r1.font.size = Pt(10.5)
    doc.add_paragraph()


def build():
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"

    h1(doc, "What 5 Years of Hospital Data Taught Me About Building Honest ML")
    subtitle(doc, "How I turned messy government data on 5,800+ US hospitals into an "
                  "explainable, deployed decision-support tool — and why I chose a "
                  "worse-looking model on purpose.")
    doc.add_paragraph()

    body(doc,
         "Every data science portfolio has the same project: a clean CSV, a model, "
         "an accuracy score, a victory lap. I wanted to build something closer to "
         "real analytics work — where the data is messy, the 'obvious' model is a "
         "trap, and the goal is a decision someone can actually act on. So I picked "
         "a domain where quality genuinely matters: US hospitals.")
    body(doc,
         "The result is Hospital Quality Intelligence — an end-to-end project built "
         "on public CMS Care Compare data that goes from SQL exploration, to two "
         "explainable machine-learning models, to a live interactive dashboard, to "
         "a five-year trend analysis. Here's the story, including the parts that broke.")

    h2(doc, "The question")
    body(doc,
         "Imagine you work for a state health department. You oversee a couple "
         "hundred hospitals, but you can only send auditors to a handful each "
         "quarter. Which ones do you inspect first? And when a hospital is "
         "struggling, what should it actually fix? Those two questions — triage and "
         "root cause — framed the entire project. They're not academic; they decide "
         "where real resources go.")

    h2(doc, "Phase 1: Learning the data with SQL")
    body(doc,
         "Before any modeling, I wrote 75 SQL queries across five perspectives — "
         "geography, ownership, clinical safety, timeliness, and patient experience. "
         "This wasn't busywork. Real CMS data is full of landmines: ratings stored "
         "as text, literal 'Not Available' strings, and three tables that only join "
         "cleanly on a facility ID. Window functions (RANK, ROW_NUMBER, NTILE, LAG), "
         "CTEs, and views turned 4,800+ hospitals into rankings, per-state leaders, "
         "and a composite risk score. The final query stitched all three datasets "
         "into a single national scorecard.")

    h2(doc, "Phase 2: The model — and the trap I almost fell into")
    body(doc,
         "I built two Random Forest models: one to predict a hospital's patient-"
         "satisfaction star rating, and one to flag likely 'underperformers' for "
         "audit triage. The classifier hit a ROC-AUC of 0.90. The regressor landed "
         "at R^2 = 0.44. But the interesting part isn't the numbers — it's a "
         "decision I made that lowered them on purpose.")
    body(doc,
         "CMS's overall star rating is mechanically derived from a set of "
         "mortality, safety, and readmission measure counts. If I fed those counts "
         "into a model predicting the rating, I'd get a near-perfect score — and "
         "prove absolutely nothing. That's data leakage: the model would just be "
         "reverse-engineering a formula, not learning anything about hospital "
         "quality.")
    pull_quote(doc,
               "\"A model that scores 0.99 by cheating is worth less than one that "
               "scores 0.90 honestly. Knowing the difference is the job.\"")
    body(doc,
         "So I excluded those columns entirely and trained only on structural and "
         "experiential features — hospital type, ownership, location, emergency "
         "services, survey engagement, and process-of-care scores. The models got "
         "weaker but honest. To make sure that discipline survived future edits, I "
         "put the forbidden columns in one place and wrote an automated test that "
         "fails if any of them ever reach the feature matrix.")

    h2(doc, "Explaining 'why', not just 'what'")
    body(doc,
         "A risk score no one understands is a risk score no one trusts. I used "
         "SHAP to open the black box and show which factors push a hospital's "
         "predicted satisfaction up or down. Survey response rate and timely-care "
         "performance carried far more weight than ownership type — a genuinely "
         "actionable finding for an administrator deciding where to invest.")
    image(doc, IMG / "shap_regression.png",
          "SHAP: what drives a hospital's predicted patient satisfaction.")

    h2(doc, "Phase 3: Five years of real trends (and a data-engineering war story)")
    body(doc,
         "A single snapshot can't tell you where a hospital is heading. So I pulled "
         "the real CMS archives — 24 quarterly releases from January 2021 to May "
         "2026 — and stacked them into a longitudinal panel of 5,830 hospitals. "
         "That let me compute each hospital's trajectory and flag 817 of them as "
         "being on a sustained decline: an early-warning list, before they even "
         "show up as low-rated.")
    body(doc,
         "Getting there was the messiest part of the whole project, and the most "
         "realistic. The 2021-2022 archives used an older CMS schema with renamed "
         "columns and categorical quality fields instead of numeric counts. Some "
         "files were UTF-8; others were Windows-1252 and crashed the parser on an "
         "en-dash. The pipeline now resolves column aliases, treats missing risk "
         "counts as genuinely unknown (NaN, not a misleading zero), and auto-detects "
         "encoding — so 24 heterogeneous releases assemble into one clean panel.")
    image(doc, IMG / "national_trend.png",
          "National quality metrics, 2021-2026, indexed to the first period.")
    body(doc,
         "The trends themselves are real and a little surprising: timely-care scores "
         "climbed through 2023, then dropped sharply in late 2024 (a CMS measure-set "
         "change, not a sudden collapse in care); patient satisfaction held steady "
         "around 3.2 stars; and survey response rates quietly eroded from 25% to "
         "23% — a small signal of nationwide survey fatigue.")

    h2(doc, "Making it real: tests, CI, and a live app")
    body(doc,
         "What separates a notebook from a product is the boring infrastructure. "
         "The project has a 13-test suite covering the cleaning logic, the leakage "
         "guard, and the trend math, all running automatically on every push via "
         "GitHub Actions across two Python versions. And the whole thing is deployed "
         "as an interactive Streamlit dashboard: pick a state, get a ranked risk "
         "list; pick a hospital, see its drivers.")

    h2(doc, "What I'd tell my past self")
    bullet(doc, "The hard part of analytics isn't the model — it's messy data and "
                "honest framing.")
    bullet(doc, "Chase the right number, not the biggest one. Guarding against "
                "leakage matters more than a leaderboard score.")
    bullet(doc, "Explainability isn't optional for decision-support. 'Why' is the "
                "product.")
    bullet(doc, "Real-world data is heterogeneous over time. Schema drift and "
                "encoding chaos are the norm, not the exception.")

    h2(doc, "Try it yourself")
    body(doc, "The full project — code, tests, and write-up — is open source, and "
              "the dashboard is live:")
    p = doc.add_paragraph()
    p.add_run("Live demo: ").bold = True
    p.add_run(DEMO)
    p = doc.add_paragraph()
    p.add_run("Code: ").bold = True
    p.add_run(REPO)
    doc.add_paragraph()
    body(doc,
         "If you're building your own portfolio project, my one piece of advice: "
         "pick a problem where being honest is harder than being impressive. That's "
         "where the real learning — and the real signal to employers — lives.")

    out = OUT_DIR / "Hospital_Quality_Intelligence_Article.docx"
    doc.save(str(out))
    print(f"Saved article -> {out}")


def build_recruiter():
    """A scannable, results-first case study aimed at recruiters and hiring managers."""
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"

    h1(doc, "Hospital Quality Intelligence")
    subtitle(doc, "Project Case Study — End-to-End Data Analytics & Machine Learning")
    body(doc,
         "An end-to-end system built on public CMS data for 5,800+ US hospitals: SQL "
         "analysis → two explainable ML models → a five-year trend analysis → a "
         "deployed, tested, CI-backed dashboard. Designed around a real decision: "
         "which hospitals should limited oversight resources target, and why.")

    kv_table(doc, [
        ("Role", "Solo project — data engineering, ML modeling, and deployment"),
        ("Data", "CMS Care Compare (public) · 3 datasets · 24 quarterly releases, 2021–2026"),
        ("Core stack", "SQL Server (T-SQL), Python, pandas, scikit-learn, SHAP, XGBoost, "
                       "Streamlit, pytest, GitHub Actions"),
        ("Live demo", DEMO),
        ("Source code", REPO),
    ])

    h2(doc, "Results at a glance")
    bullet(doc, "75 SQL queries across 5 analytical perspectives on 5,830 hospitals")
    bullet(doc, "Underperformer-triage classifier: ROC-AUC 0.90, 84% recall")
    bullet(doc, "Patient-satisfaction regressor: R^2 0.44 (leakage-free, cross-validated)")
    bullet(doc, "24 quarterly releases stitched into a longitudinal panel; 817 hospitals "
                "flagged on a declining trajectory")
    bullet(doc, "13 automated tests, CI on Python 3.10 & 3.11, and a publicly deployed app")

    h2(doc, "Skills demonstrated")
    kv_table(doc, [
        ("Data Engineering", "Multi-source ETL; schema-drift handling across 5 years; "
                             "encoding normalization; long→wide reshaping; 100k+ row panel"),
        ("SQL", "Joins, CTEs, window functions (RANK/NTILE/LAG), views, composite scoring"),
        ("Machine Learning", "Classification + regression; 5-fold cross-validation; "
                             "baseline-anchored model selection; data-leakage prevention"),
        ("Explainable AI", "SHAP feature attribution for stakeholder-facing 'why'"),
        ("Software / MLOps", "pytest test suite; GitHub Actions CI; config-driven design; "
                             "Streamlit Cloud deployment"),
        ("Communication", "Interactive dashboard, professional documentation, and "
                          "business-framed problem definition"),
    ])

    h2(doc, "What I built")
    body(doc, "Phase 1 — SQL Analytics: 75 queries turning three messy CMS tables into "
              "rankings, per-state leaders, and a national hospital scorecard.")
    body(doc, "Phase 2 — Machine Learning: two Random Forest models (satisfaction "
              "prediction + underperformer triage), selected via cross-validated "
              "benchmarking and explained with SHAP. Deployed as a 3-tab Streamlit "
              "dashboard: ranked audit list, hospital drill-down, and model card.")
    body(doc, "Phase 3 — Longitudinal Pipeline: 24 real CMS quarterly releases "
              "(2021–2026) assembled into a panel that flags hospitals on a sustained "
              "downward trajectory — an early-warning tool, not just a snapshot.")

    h2(doc, "Why it stands out")
    bullet(doc, "Judgment over vanity metrics — identified and prevented data leakage, "
                "choosing an honest 0.90 model over a fake 'perfect' one, enforced by a test.")
    bullet(doc, "Real-world resilience — handled genuine messy government data: "
                "multi-year schema changes and mixed file encodings, not a clean Kaggle CSV.")
    bullet(doc, "Production mindset — tested, CI-verified, documented, and deployed live, "
                "not a one-off notebook.")

    image(doc, IMG / "national_trend.png",
          "Real national quality trends, 2021–2026 — one output of the longitudinal pipeline.")

    h2(doc, "Explore the work")
    p = doc.add_paragraph()
    p.add_run("Live demo: ").bold = True
    p.add_run(DEMO)
    p = doc.add_paragraph()
    p.add_run("Code & documentation: ").bold = True
    p.add_run(REPO)

    out = OUT_DIR / "Hospital_Quality_Intelligence_Recruiter.docx"
    doc.save(str(out))
    print(f"Saved recruiter case study -> {out}")


if __name__ == "__main__":
    build()
    build_recruiter()
